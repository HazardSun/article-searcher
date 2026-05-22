"""
文件扫描与解析模块
支持 Markdown (.md) 和 HTML (.html) 文件的递归扫描与内容提取
"""

import os
import io
import hashlib
from pathlib import Path
from dataclasses import dataclass, field
from typing import List, Optional
from datetime import datetime

from bs4 import BeautifulSoup
import marko


@dataclass
class FileMetadata:
    """文件元数据"""
    file_path: str
    file_name: str
    file_extension: str
    file_size: int
    md5_hash: str
    modified_time: float
    title: str = ""
    tags: List[str] = field(default_factory=list)


@dataclass
class TextChunk:
    """文本切片"""
    chunk_id: str
    file_path: str
    file_name: str
    title: str
    content: str
    start_line: int
    end_line: int
    chunk_index: int
    total_chunks: int
    metadata: dict = field(default_factory=dict)


class FileScanner:
    """文件扫描器 - 递归扫描指定目录中的 MD、HTML、TXT、PDF、DOCX 文件"""

    SUPPORTED_EXTENSIONS = {'.md', '.html', '.htm', '.txt', '.pdf', '.docx'}

    def __init__(self):
        self.scanned_files: List[FileMetadata] = []

    def scan_directory(self, directory: str) -> List[FileMetadata]:
        """递归扫描目录，返回所有支持格式的文件元数据"""
        self.scanned_files = []
        root_path = Path(directory)

        if not root_path.exists():
            raise FileNotFoundError(f"目录不存在: {directory}")

        for ext in self.SUPPORTED_EXTENSIONS:
            for file_path in root_path.rglob(f'*{ext}'):
                if file_path.is_file():
                    meta = self._extract_metadata(file_path)
                    self.scanned_files.append(meta)

        return self.scanned_files

    def _extract_metadata(self, file_path: Path) -> FileMetadata:
        """提取单个文件的元数据"""
        stat = file_path.stat()
        md5 = self._compute_md5(file_path)

        ext = file_path.suffix
        if ext in ('.pdf', '.docx'):
            title = self._extract_title_binary(file_path, ext)
            content = ""
        else:
            content = file_path.read_text(encoding='utf-8', errors='ignore')
            title = self._extract_title(content, ext)

        return FileMetadata(
            file_path=str(file_path),
            file_name=file_path.name,
            file_extension=ext,
            file_size=stat.st_size,
            md5_hash=md5,
            modified_time=stat.st_mtime,
            title=title
        )

    @staticmethod
    def _compute_md5(file_path: Path) -> str:
        """计算文件 MD5 哈希值"""
        hasher = hashlib.md5()
        with open(file_path, 'rb') as f:
            for chunk in iter(lambda: f.read(8192), b''):
                hasher.update(chunk)
        return hasher.hexdigest()

    @staticmethod
    def _extract_title(content: str, extension: str) -> str:
        """从文本文件内容中提取标题"""
        if extension in ('.md',):
            lines = content.split('\n')
            for line in lines[:20]:
                stripped = line.strip()
                if stripped.startswith('# '):
                    return stripped[2:].strip()
        elif extension in ('.html', '.htm'):
            soup = BeautifulSoup(content, 'lxml')
            title_tag = soup.find('title')
            if title_tag and title_tag.get_text(strip=True):
                return title_tag.get_text(strip=True)
            h1 = soup.find('h1')
            if h1 and h1.get_text(strip=True):
                return h1.get_text(strip=True)
        elif extension == '.txt':
            lines = content.split('\n')
            for line in lines[:10]:
                stripped = line.strip()
                if stripped:
                    return stripped[:80]

        return ""

    @staticmethod
    def _extract_title_binary(file_path: Path, extension: str) -> str:
        """从二进制文件（PDF/DOCX）中提取标题"""
        try:
            if extension == '.pdf':
                import fitz
                doc = fitz.open(str(file_path))
                for page in doc:
                    text = page.get_text().strip()
                    if text:
                        title = text.split('\n')[0].strip()[:80]
                        doc.close()
                        return title
                doc.close()
            elif extension == '.docx':
                import docx
                d = docx.Document(str(file_path))
                for para in d.paragraphs:
                    if para.text.strip():
                        return para.text.strip()[:80]
        except Exception:
            pass
        return ""


class ContentParser:
    """内容解析器 - 将 MD/HTML/TXT/PDF/DOCX 解析为结构化段落"""

    def parse_markdown(self, content: str) -> List[dict]:
        """
        解析 Markdown 内容为结构化段落
        返回: [{"type": "heading"|"paragraph"|"code", "content": str, "line": int}]
        """
        paragraphs = []
        lines = content.split('\n')
        current_paragraph = []
        current_start_line = 0
        in_code_block = False
        code_content = []
        code_start_line = 0

        for i, line in enumerate(lines):
            stripped = line.strip()

            if stripped.startswith('```'):
                if in_code_block:
                    code_content.append(stripped)
                    paragraphs.append({
                        'type': 'code',
                        'content': '\n'.join(code_content),
                        'line': code_start_line
                    })
                    code_content = []
                    in_code_block = False
                else:
                    if current_paragraph:
                        paragraphs.append({
                            'type': 'paragraph',
                            'content': '\n'.join(current_paragraph),
                            'line': current_start_line
                        })
                        current_paragraph = []
                    in_code_block = True
                    code_start_line = i
                    code_content = [stripped]
                continue

            if in_code_block:
                code_content.append(line)
                continue

            if stripped.startswith('#'):
                if current_paragraph:
                    paragraphs.append({
                        'type': 'paragraph',
                        'content': '\n'.join(current_paragraph),
                        'line': current_start_line
                    })
                    current_paragraph = []
                paragraphs.append({
                    'type': 'heading',
                    'content': stripped.lstrip('#').strip(),
                    'line': i,
                    'level': len(stripped) - len(stripped.lstrip('#'))
                })
                continue

            if stripped.startswith(('-', '*', '+', '1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.', '0.')):
                if current_paragraph:
                    paragraphs.append({
                        'type': 'paragraph',
                        'content': '\n'.join(current_paragraph),
                        'line': current_start_line
                    })
                    current_paragraph = []
                paragraphs.append({
                    'type': 'list',
                    'content': stripped,
                    'line': i
                })
                continue

            if stripped == '':
                if current_paragraph:
                    paragraphs.append({
                        'type': 'paragraph',
                        'content': '\n'.join(current_paragraph),
                        'line': current_start_line
                    })
                    current_paragraph = []
                continue

            if not current_paragraph:
                current_start_line = i
            current_paragraph.append(line)

        if current_paragraph:
            paragraphs.append({
                'type': 'paragraph',
                'content': '\n'.join(current_paragraph),
                'line': current_start_line
            })

        return paragraphs

    def parse_html(self, content: str) -> List[dict]:
        """
        解析 HTML 内容为结构化段落
        返回: [{"type": "heading"|"paragraph"|"code", "content": str, "line": int}]
        """
        soup = BeautifulSoup(content, 'lxml')

        for tag in soup.find_all(['script', 'style', 'nav', 'footer', 'header']):
            tag.decompose()

        paragraphs = []
        line_counter = 0

        for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'li', 'pre', 'div', 'article', 'section']):
            text = element.get_text(separator=' ', strip=True)
            if not text:
                continue

            tag_name = element.name
            if tag_name.startswith('h'):
                paragraphs.append({
                    'type': 'heading',
                    'content': text,
                    'line': line_counter,
                    'level': int(tag_name[1])
                })
            elif tag_name == 'pre':
                paragraphs.append({
                    'type': 'code',
                    'content': text,
                    'line': line_counter
                })
            elif tag_name in ('p', 'li', 'div', 'article', 'section'):
                if len(text) > 10:
                    paragraphs.append({
                        'type': 'paragraph',
                        'content': text,
                        'line': line_counter
                    })

            line_counter += 1

        if not paragraphs:
            body_text = soup.get_text(separator='\n', strip=True)
            if body_text:
                for i, line in enumerate(body_text.split('\n')):
                    if line.strip():
                        paragraphs.append({
                            'type': 'paragraph',
                            'content': line.strip(),
                            'line': i
                        })

        return paragraphs

    @staticmethod
    def parse_text(content: str) -> List[dict]:
        """解析纯文本内容为段落"""
        paragraphs = []
        current = []
        start_line = 0
        lines = content.split('\n')

        for i, line in enumerate(lines):
            stripped = line.strip()
            if stripped == '':
                if current:
                    paragraphs.append({
                        'type': 'paragraph',
                        'content': '\n'.join(current),
                        'line': start_line
                    })
                    current = []
                continue

            if not current:
                start_line = i
            current.append(stripped)

        if current:
            paragraphs.append({
                'type': 'paragraph',
                'content': '\n'.join(current),
                'line': start_line
            })

        return paragraphs

    @staticmethod
    def parse_pdf(file_path: str) -> List[dict]:
        """解析 PDF 文件为段落列表"""
        paragraphs = []
        try:
            import fitz
            doc = fitz.open(file_path)
            line_counter = 0
            for page in doc:
                text = page.get_text().strip()
                if not text:
                    continue
                for block in page.get_text("blocks"):
                    block_text = block[4].strip() if len(block) > 4 else ''
                    if not block_text:
                        continue
                    paragraphs.append({
                        'type': 'paragraph',
                        'content': block_text,
                        'line': line_counter
                    })
                    line_counter += 1
            doc.close()
        except Exception as e:
            import logging
            logging.getLogger(__name__).error(f"PDF parse failed: {e}")

        return paragraphs

    @staticmethod
    def parse_docx(file_path: str) -> List[dict]:
        """解析 DOCX 文件为段落列表"""
        paragraphs = []
        try:
            import docx
            d = docx.Document(file_path)
            for i, para in enumerate(d.paragraphs):
                text = para.text.strip()
                if not text:
                    continue
                paragraphs.append({
                    'type': 'paragraph',
                    'content': text,
                    'line': i
                })
        except Exception as e:
            import logging
            logging.getLogger(__name__).error(f"DOCX parse failed: {e}")

        return paragraphs
